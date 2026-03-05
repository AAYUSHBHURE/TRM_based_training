from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_functional_spec():
    doc = Document()

    # Style helper
    def add_heading(text, level):
        h = doc.add_heading(text, level=level)
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    # Title
    title = doc.add_heading('Sprint Functional Document for TRM Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Introduction
    add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "The Tiny Recursive Model (TRM) project aims to revolutionize efficient AI reasoning "
        "by challenging the 'depth via parameters' paradigm. Sprint 1 is focused on implementing "
        "the core recursive engine and validating it on Sudoku-Extreme puzzles, proving that "
        "small models can solve complex logic through iteration."
    )

    # 2. Product Goal
    add_heading('2. Product Goal', level=1)
    doc.add_paragraph(
        "The primary goal of this sprint is to achieve >87% accuracy on Sudoku tasks "
        "using a model with fewer than 7M parameters. This aligns with the overarching objective "
        "of enabling powerful reasoning capabilities on edge devices by trading compute time for model size."
    )

    # 3. Demography
    add_heading('3. Demography (Users, Location)', level=1)
    
    p = doc.add_paragraph()
    p.add_run("Users").bold = True
    doc.add_paragraph("Target Users: AI Researchers, Edge Computing Engineers, Educational Tech Students")
    doc.add_paragraph("User Characteristics: Technical proficiency in Python/PyTorch, interest in model efficiency")
    
    p = doc.add_paragraph()
    p.add_run("Location").bold = True
    doc.add_paragraph("Target Location: Global Open Source Community & Research Labs")

    # 4. Business Processes
    add_heading('4. Business Processes', level=1)
    doc.add_paragraph("Key business processes include:")
    
    p = doc.add_paragraph()
    p.add_run("Data Generation:").bold = True
    doc.add_paragraph("Process for procedurally generating valid, unique Sudoku grids and Mazes with guaranteed solutions.")
    
    p = doc.add_paragraph()
    p.add_run("Model Training:").bold = True
    doc.add_paragraph("Process for training the recursive block using Deep Supervision to ensure every cycle improves the state.")
    
    p = doc.add_paragraph()
    p.add_run("Inference & Reasoning:").bold = True
    doc.add_paragraph("Process where the model consumes a puzzle and iteratively refines its internal state over T cycles to produce a solution.")

    # 5. Features
    add_heading('5. Features', level=1)
    doc.add_paragraph("This sprint will focus on implementing the following key features:")

    # Feature 1
    add_heading('5.1 Recursive Inference Engine', level=2)
    p = doc.add_paragraph()
    p.add_run("1. Description").bold = True
    doc.add_paragraph(
        "A lightweight Transformer block that loops its output back as input. "
        "It maintains a persistent latent state 'z' that represents the current understanding of the puzzle."
    )
    p = doc.add_paragraph()
    p.add_run("2. User Story").bold = True
    doc.add_paragraph(
        "As a researcher, I want to configure the number of recurrence cycles (T) at inference time "
        "so that I can trade off speed for accuracy without retraining."
    )

    # Feature 2
    add_heading('5.2 Deep Supervision Loss', level=2)
    p = doc.add_paragraph()
    p.add_run("1. Description").bold = True
    doc.add_paragraph(
        "A training mechanism that calculates loss at every step of the recursion, "
        "forcing the model to learn a stable convergence trajectory."
    )
    p = doc.add_paragraph()
    p.add_run("2. User Story").bold = True
    doc.add_paragraph(
        "As a developer, I want the model to receive feedback on intermediate thinking steps "
        "so that it minimizes the 'vanishing gradient' problem during training."
    )

    # 6. Authorization Matrix
    add_heading('6. Authorization Matrix', level=1)
    doc.add_paragraph("Define the roles and their corresponding access levels:")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Access Level'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    roles = [
        ('Researcher', 'Full access to model architecture and training hyperparameters'),
        ('Developer', 'Access to code, data generation scripts, and evaluation metrics'),
        ('Demo User', 'Read-only access to run inference on puzzles and view results')
    ]
    
    for role, access in roles:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = access

    # 7. Assumptions
    add_heading('7. Assumptions', level=1)
    doc.add_paragraph("The development environment and infrastructure will remain stable during the sprint.")
    doc.add_paragraph("The hardware available (consumer GPUs) is sufficient for training small models.")
    doc.add_paragraph("Team members possess the necessary skills to implement recursive transformers.")

    # Save
    doc.save('Functional Specification.docx')
    print("Functional Specification.docx created successfully.")

if __name__ == "__main__":
    create_functional_spec()
